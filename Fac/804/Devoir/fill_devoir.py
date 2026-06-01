import docx
import shutil

# Backup
shutil.copy('Devoir1_INFO0804.docx', 'Devoir1_INFO0804_backup.docx')

doc = docx.Document('Devoir1_INFO0804.docx')

# --- helpers ---

def set_cell(table, row, col, text):
    cell = table.rows[row].cells[col]
    para = cell.paragraphs[0]
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = str(text)
    else:
        para.add_run(str(text))

def set_para(para, text):
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)

def set_code_cell(table_idx, lines):
    table = doc.tables[table_idx]
    cell = table.rows[0].cells[0]
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ''
    if cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = lines[0]
    else:
        cell.paragraphs[0].add_run(lines[0])
    for line in lines[1:]:
        p = cell.add_paragraph()
        r = p.add_run(line)
        r.font.name = 'Courier New'


# --- TABLE 2 - Partie I : Test/Train splits ---
t2 = doc.tables[2]
data_I = [
    (1, [0.4942, 0.7763, 1.0853, 0.6561]),
    (2, [0.4920, 0.7451, 0.8676, 0.6650]),
    (3, [0.4286, 0.7826, 0.8531, 0.7007]),
    (4, [0.3736, 0.8203, 0.7756, 0.7188]),
    (5, [0.4309, 0.7970, 0.5945, 0.7576]),
    (6, [0.4131, 0.8062, 0.5711, 0.7792]),
    (7, [0.4052, 0.8061, 0.5451, 0.7532]),
]
for row, vals in data_I:
    for ci, v in enumerate(vals):
        set_cell(t2, row, ci+2, f'{v:.4f}')

# --- TABLE 3 - Code Partie II ModelCheckpoint ---
checkpoint_code = [
    'from tensorflow.keras.callbacks import ModelCheckpoint',
    '',
    'checkpoint = ModelCheckpoint(',
    "    'meilleur_modele.keras',",
    "    monitor='val_loss',",
    '    save_best_only=True,',
    '    verbose=0',
    ')',
    '',
    'class_model = model.fit(',
    '    X_train, y_train,',
    '    validation_data=(X_test, y_test),',
    '    epochs=500, batch_size=10,',
    '    callbacks=[checkpoint]',
    ')',
    '',
    '# => Meilleur modele a l epoch 360 (val_loss=0.5255)',
]
set_code_cell(3, checkpoint_code)

# --- TABLE 4 - Partie III : Epochs ---
t4 = doc.tables[4]
data_III = [
    (1, [0.6865, 0.6369, 0.6431, 0.6970]),
    (2, [0.5891, 0.6946, 0.5821, 0.7229]),
    (3, [0.5315, 0.7169, 0.5665, 0.7749]),
    (4, [0.4908, 0.7505, 0.5992, 0.7489]),
    (5, [0.4800, 0.7449, 0.6408, 0.7359]),
    (6, [0.4613, 0.7747, 0.5695, 0.7186]),
    (7, [0.4325, 0.8007, 0.6175, 0.7229]),
    (8, [0.4409, 0.7616, 0.6742, 0.7359]),
    (9, [0.4175, 0.8082, 0.7875, 0.7359]),
]
for row, vals in data_III:
    for ci, v in enumerate(vals):
        set_cell(t4, row, ci+1, f'{v:.4f}')

# --- TABLE 5 - Partie IV : nb_batchs ---
t5 = doc.tables[5]
set_cell(t5, 1, 0, 'nb_batchs/epoch')
batch_counts = ['54/54', '34/34', '17/17', '9/9', '5/5', '1/1']
for ci, val in enumerate(batch_counts):
    set_cell(t5, 1, ci+3, val)

# --- TABLE 7 - Partie V : Batch sizes performances ---
t7 = doc.tables[7]
data_V = [
    (1, [0.4718, 0.7616, 0.6226, 0.7403]),
    (2, [0.4985, 0.7412, 0.5684, 0.7359]),
    (3, [0.5415, 0.7374, 0.5841, 0.7619]),
    (4, [0.5515, 0.7281, 0.5882, 0.7143]),
    (5, [0.5811, 0.6853, 0.6306, 0.7186]),
    (6, [0.5739, 0.7002, 0.5913, 0.6926]),
    (7, [0.6095, 0.6834, 0.6245, 0.6883]),
    (8, [0.8340, 0.6387, 0.8041, 0.7100]),
]
for row, vals in data_V:
    for ci, v in enumerate(vals):
        set_cell(t7, row, ci+1, f'{v:.4f}')

# --- TABLE 12 - Code Partie VII ---
code_VII = [
    '# Etape 1 : 30% pour le test',
    'X_train_full, X_test, y_train_full, y_test = train_test_split(',
    '    X, Y, test_size=0.30, random_state=seed)',
    '',
    '# Etape 2 : 15% du reste pour la validation (~10% de la base totale)',
    'X_train, X_val, y_train, y_val = train_test_split(',
    '    X_train_full, y_train_full, test_size=0.15, random_state=seed)',
    '# => Train=456 | Validation=81 | Test=231',
    '',
    '# Code pour retourner le meilleur modele',
    'checkpoint = ModelCheckpoint(',
    "    'meilleur_modele_VII.keras',",
    "    monitor='val_loss', save_best_only=True, verbose=0)",
    '',
    'class_model = model.fit(',
    '    X_train, y_train,',
    '    validation_data=(X_val, y_val),',
    '    epochs=500, batch_size=10,',
    '    callbacks=[checkpoint])',
    '',
    '# Modele final  : loss=0.7494, accuracy=0.7100',
    '# Meilleur modele: loss=0.7793, accuracy=0.6970',
]
set_code_cell(12, code_VII)

# --- TABLE 13 - Code 4-neuron model ---
code_4n = [
    'from tensorflow.keras import initializers',
    'from tensorflow.keras.models import Sequential',
    'from tensorflow.keras.layers import Dense',
    '',
    'seed = 5',
    'initializer = initializers.GlorotUniform(seed=seed)',
    '',
    'model = Sequential()',
    'model.add(Dense(4, input_shape=(8,),',
    '               kernel_initializer=initializer,',
    "               activation='relu'))",
    'model.add(Dense(1, kernel_initializer=initializer,',
    "               activation='sigmoid'))",
    "model.compile(loss='binary_crossentropy',",
    "              optimizer='adam', metrics=['accuracy'])",
]
set_code_cell(13, code_4n)

# --- TABLE 15 - Architectures classification report ---
t15 = doc.tables[15]
arch_data = [
    (1, ['0.48', '0.68', '0.68', '0.56']),
    (2, ['0.73', '0.68', '0.68', '0.69']),
    (3, ['0.67', '0.71', '0.71', '0.66']),
]
for row, vals in arch_data:
    for ci, v in enumerate(vals):
        set_cell(t15, row, ci+1, v)

# --- TABLE 16 - Code change neurones ---
code_neurons = [
    '# Changer le nombre de neurones dans la couche cachee',
    'n_neurons = 8  # modifier : 4, 8, 16, ...',
    '',
    'model = Sequential()',
    'model.add(Dense(n_neurons, input_shape=(8,),',
    '               kernel_initializer=initializer,',
    "               activation='relu'))",
    'model.add(Dense(1, kernel_initializer=initializer,',
    "               activation='sigmoid'))",
    "model.compile(loss='binary_crossentropy',",
    "              optimizer='adam', metrics=['accuracy'])",
]
set_code_cell(16, code_neurons)

# --- TABLE 17 - Neurons variation ---
t17 = doc.tables[17]
neuron_data = [
    (1, ['0.48', '0.68', '0.68', '0.56']),
    (2, ['0.78', '0.77', '0.77', '0.78']),
    (3, ['0.77', '0.77', '0.77', '0.77']),
]
for row, vals in neuron_data:
    for ci, v in enumerate(vals):
        set_cell(t17, row, ci+1, v)

# --- PARAGRAPHS - Text answers ---
paras = doc.paragraphs

set_para(paras[11],
    "Observations : Plus la taille d'apprentissage est grande, plus la val_loss diminue et "
    "la val_accuracy augmente. Avec 10% d'apprentissage (val_acc=0.6561), le modele "
    "generalise mal. Le meilleur compromis est obtenu autour de 80% d'apprentissage (val_acc=0.7792).")

set_para(paras[12],
    "Conclusions : Il existe un compromis optimal autour de 70-80% pour l'apprentissage. "
    "En dessous de 50%, le modele sous-apprend faute de donnees suffisantes. "
    "La val_loss diminue systematiquement quand la taille d'apprentissage augmente.")

set_para(paras[15],
    "On observe du sur-apprentissage (overfitting) : la val_loss atteint son minimum vers "
    "l'epoch 360 (val_loss=0.5255) puis remonte progressivement, tandis que la train_loss "
    "continue de diminuer. Le modele memorise les donnees d'entrainement au detriment de la generalisation.")

set_para(paras[21],
    "Observations : Les meilleures performances de validation sont obtenues a 100 epochs "
    "(val_acc=0.7749, val_loss=0.5665). Au-dela, la val_loss augmente progressivement "
    "indiquant un sur-apprentissage, tandis que la train_acc continue de croitre.")

set_para(paras[22],
    "Conclusions : Un trop grand nombre d'epochs provoque du sur-apprentissage. "
    "Il est recommande d'utiliser ModelCheckpoint (save_best_only=True) pour conserver "
    "le meilleur modele. Environ 100 epochs semble optimal pour ce jeu de donnees.")

set_para(paras[36],
    "537/537 : le numerateur indique le numero du batch courant traite, le denominateur "
    "indique le nombre total de batchs par epoch. "
    "Avec N=537 exemples et batch_size=1 : ceil(537/1)=537 batchs par epoch. "
    "Formule : nb_batchs = ceil(N / K) avec N=taille base, K=taille du batch.")

set_para(paras[39],
    "La valeur Max possible est N = 537 (taille totale de l'ensemble d'apprentissage). "
    "Avec batch_size=N, un seul batch contient tous les exemples, ce qui donne "
    "1 seule mise a jour des poids par epoch (batch gradient descent complet). "
    "Au-dela de N il n'y a plus d'exemples supplementaires a inclure.")

set_para(paras[42],
    "Observations : batch_size=1 est le plus lent (121s pour 50 epochs). "
    "Les batchs intermediaires (54-160) sont les plus rapides (~15-18s). "
    "Les tres grands batchs restent plus lents que les batchs moyens.")

set_para(paras[43],
    "Conclusions : Un batch_size trop petit (=1) augmente considerablement le temps de "
    "calcul a cause du grand nombre de mises a jour des poids. Un batch_size modere "
    "(10-64) offre le meilleur compromis entre vitesse de calcul et qualite du gradient.")

set_para(paras[49],
    "Observations : Les petits batchs (1, 2, 10) donnent une meilleure val_accuracy grace "
    "au bruit stochastique du gradient (effet regularisant). Les grands batchs (128, Max=537) "
    "convergent moins bien (val_acc=0.6883 et 0.7100) avec train_loss plus elevee.")

set_para(paras[50],
    "Conclusions : Un batch_size de 10 offre le meilleur compromis (val_acc=0.7619). "
    "Les tres grands batchs convergent vers de moins bonnes solutions car le gradient "
    "est trop lisse, ce qui favorise des minima locaux moins generalisant.")

set_para(paras[69],
    "Observations : Sans validation_data, model.evaluate donne loss=0.5805, accuracy=0.7662. "
    "La matrice de confusion : [[145,15],[39,32]] montre 145 vrais negatifs, 32 vrais positifs, "
    "15 faux positifs et 39 faux negatifs. Le modele predit mieux la classe 0 (non-diabetique).")

set_para(paras[70],
    "Conclusions : Entrainer sans validation_data empeche de detecter le sur-apprentissage "
    "pendant l'entrainement. Il est preferable d'utiliser validation_data pour surveiller "
    "la generalisation et arreter l'entrainement au bon moment avec ModelCheckpoint.")

set_para(paras[87],
    "model.get_config() retourne un dictionnaire JSON decrivant l'architecture complete : "
    "le type de modele (Sequential), la liste des couches avec leur type (Dense), "
    "le nombre de neurones (units), la fonction d'activation, l'initialiseur des poids, "
    "le type de regularisation, etc. Utile pour sauvegarder et reconstruire l'architecture.")

set_para(paras[89],
    "model.summary() affiche un tableau recapitulatif : "
    "dense (Dense) | output=(None,4) | params=36 ; "
    "dense_1 (Dense) | output=(None,1) | params=5. "
    "Total params: 41 (164 B). Trainable: 41, Non-trainable: 0.")

set_para(paras[91],
    "print(model.get_weights()) retourne une liste de 4 tableaux numpy : "
    "weight[0] shape=(8,4) = poids couche 1 (8 entrees -> 4 neurones), "
    "weight[1] shape=(4,) = biais couche 1 (4 valeurs), "
    "weight[2] shape=(4,1) = poids couche sortie (4 -> 1), "
    "weight[3] shape=(1,) = biais sortie. Total: 32+4+4+1=41 parametres.")

set_para(paras[103],
    "Nombre total = (8x4)+4 + (4x1)+1 = 32+4+4+1 = 41 parametres. "
    "Couche cachee : 8 entrees x 4 neurones = 32 poids + 4 biais = 36. "
    "Couche sortie : 4 entrees x 1 neurone = 4 poids + 1 biais = 5. "
    "Total = 41, confirme par model.summary() (Total params: 41).")

set_para(paras[106],
    "Observations : La configuration 1 couche de 4 neurones predit uniquement la classe 0 "
    "(precision classe 1 = 0.00) - reseau trop simple. "
    "3 couches (4,5,6) et 4 couches (4,8,3) donnent de meilleures performances "
    "equilibrees entre les deux classes (accuracy 68-71%).")

set_para(paras[107],
    "Conclusions : La profondeur seule n'ameliore pas toujours les performances. "
    "Un reseau trop simple (4 neurones) sous-apprend. L'architecture (4,8,3) avec 71% "
    "d'accuracy offre le meilleur equilibre sur ce jeu de donnees.")

set_para(paras[112],
    "Observations : Avec 4 neurones, le modele ne predit que la classe 0 (accuracy=0.68). "
    "Avec 8 neurones, l'accuracy monte a 77% et la classe 1 est bien reconnue (f1=0.65). "
    "16 neurones donne des resultats similaires a 8 (accuracy=0.77).")

set_para(paras[113],
    "Conclusions : Le nombre de neurones est crucial. Trop peu (4) ne permet pas d'apprendre "
    "la frontiere de decision non-lineaire. 8 ou 16 neurones donnent des resultats optimaux. "
    "Au-dela, le gain est marginal et le risque de sur-apprentissage augmente.")

doc.save('Devoir1_INFO0804.docx')
print("Docx rempli avec succes !")
