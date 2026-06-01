# -*- coding: utf-8 -*-
"""Remplissage automatique des réponses dans Devoir2_INFO0804.docx"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import copy

PLACEHOLDER = '\u2026'  # …

def set_run_text(run, text):
    """Remplace le texte d'un run en conservant la mise en forme."""
    run.text = text

def fill_question_run(paragraph, answer_text):
    """Remplace le placeholder '…' dans run[2] par la réponse."""
    run = paragraph.runs[2]
    original = run.text
    # Remplacer les '…' répétés à la fin par la réponse
    idx = original.find(PLACEHOLDER)
    if idx != -1:
        run.text = original[:idx] + answer_text


def add_text_to_empty_paragraph(paragraph, text, bold=False, italic=False):
    """Ajoute un run avec du texte dans un paragraphe vide."""
    run = paragraph.add_run(text)
    run.font.name = 'Times'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic


def replace_placeholder_paragraph(paragraph, text):
    """Remplace le contenu '…' d'un paragraphe par le texte fourni."""
    run = paragraph.runs[0]
    run.text = text


doc = Document('Devoir2_INFO0804.docx')
paragraphs = doc.paragraphs

# ─────────────────────────────────────────────────────────────────────────────
# EXO 1 – Régression linéaire (Housing)
# ─────────────────────────────────────────────────────────────────────────────

# I – Initialisation des poids (paragraph 18)
fill_question_run(
    paragraphs[18],
    "La méthode optimale est uniform car sur le test son mean_test_score est de "
    "-33.252957 (std: 5.765042), ce qui est le score le plus élevé (le moins négatif) "
    "parmi toutes les méthodes d'initialisation testées."
)

# II – Fonction d'activation (paragraph 29)
fill_question_run(
    paragraphs[29],
    "La méthode optimale est softplus car sur le test son mean_test_score est de "
    "-28.958168 (std: 3.521216), ce qui est le meilleur score parmi toutes les "
    "fonctions d'activation testées."
)

# III – Fonction coût (paragraph 33)
fill_question_run(
    paragraphs[33],
    "La méthode optimale est mean_squared_error car sur le test son mean_test_score "
    "est de -25.258120 (std: 4.223791), ce qui représente la meilleure performance "
    "parmi les fonctions coût testées."
)

# IV – Optimizer (paragraph 40)
fill_question_run(
    paragraphs[40],
    "La méthode optimale est RMSprop car sur le test son mean_test_score est de "
    "-30.021903 (std: 9.163522), ce qui est le meilleur score parmi tous les "
    "optimiseurs testés."
)

# V – Dropout / weight_constraint (paragraph 46)
fill_question_run(
    paragraphs[46],
    "La combinaison optimale est Dropout=0.0 et weight_constraint=2 car sur le test "
    "son mean_test_score est de -28.652179 (std: 3.866792), ce qui est le meilleur "
    "score parmi toutes les combinaisons testées."
)

# Construire un modèle optimal Exo1 (paragraph 49 – vide)
code_exo1 = (
    "from keras.models import Sequential\n"
    "from keras.layers import Dense, Dropout\n"
    "from keras.constraints import max_norm\n"
    "from scikeras.wrappers import KerasRegressor\n\n"
    "def create_optimal_model():\n"
    "    model = Sequential()\n"
    "    model.add(Dense(13, input_dim=13,\n"
    "                    kernel_initializer='uniform',\n"
    "                    activation='softplus',\n"
    "                    kernel_constraint=max_norm(2)))\n"
    "    model.add(Dropout(0.0))\n"
    "    model.add(Dense(1))\n"
    "    model.compile(loss='mean_squared_error', optimizer='RMSprop')\n"
    "    return model\n\n"
    "optimal_model = KerasRegressor(build_fn=create_optimal_model,\n"
    "                               epochs=100, batch_size=10, verbose=1)\n"
    "history = optimal_model.fit(X_train, y_train,\n"
    "                            validation_data=(X_test, y_test))\n"
    "# Résultats : Train_Loss=17.6971 | Train_MAE=3.1349\n"
    "#             Val_Loss=47.4620   | Val_MAE=5.5789"
)
add_text_to_empty_paragraph(paragraphs[49], code_exo1)

# Observations Exo1 (paragraph 51)
obs_exo1 = (
    "Le modèle optimal obtient une perte MSE de 17.6971 sur les données "
    "d'entraînement avec un MAE de 3.1349. Sur les données de validation, "
    "la perte MSE monte à 47.4620 et le MAE à 5.5789. On observe un écart "
    "notable entre les performances d'entraînement et de validation, ce qui "
    "indique une tendance au surapprentissage (overfitting). L'utilisation "
    "de softplus comme fonction d'activation et de RMSprop comme optimiseur "
    "semble bien adaptée à ce problème de régression. Le dropout nul et le "
    "weight_constraint=2 n'ont pas suffi à éliminer complètement l'overfitting, "
    "suggérant qu'une augmentation de la régularisation pourrait améliorer "
    "la généralisation."
)
replace_placeholder_paragraph(paragraphs[51], obs_exo1)

# Conclusions Exo1 (paragraph 53)
conc_exo1 = (
    "La recherche par grille a permis d'identifier les hyperparamètres optimaux "
    "pour la régression sur la base Housing : initialisation uniform, activation "
    "softplus, fonction coût mean_squared_error, optimiseur RMSprop, dropout=0.0 "
    "et weight_constraint=2. Le modèle final atteint un MAE de 3.1349 en "
    "entraînement et 5.5789 en validation. L'écart train/validation indique un "
    "certain surapprentissage ; une architecture plus profonde ou une régularisation "
    "renforcée pourrait améliorer les performances en généralisation. La méthode de "
    "tuning par GridSearchCV s'est révélée efficace pour identifier automatiquement "
    "les meilleures combinaisons d'hyperparamètres."
)
replace_placeholder_paragraph(paragraphs[53], conc_exo1)

# ─────────────────────────────────────────────────────────────────────────────
# EXO 2 – Classification binaire (Diabetes)
# ─────────────────────────────────────────────────────────────────────────────

# I – Initialisation des poids (paragraph 63)
fill_question_run(
    paragraphs[63],
    "La méthode optimale est lecun_uniform car sur le test son mean_test_score est "
    "de 0.719808 (std: 0.047222), ce qui est le score le plus élevé parmi toutes "
    "les méthodes d'initialisation testées."
)

# II – Fonction d'activation (paragraph 70)
fill_question_run(
    paragraphs[70],
    "La méthode optimale est softplus car sur le test son mean_test_score est de "
    "0.710028 (std: 0.034506), ce qui est le score le plus élevé parmi toutes les "
    "fonctions d'activation testées."
)

# III – Fonction coût (paragraph 75)
fill_question_run(
    paragraphs[75],
    "La méthode optimale est binary_crossentropy car sur le test son mean_test_score "
    "est de 0.696405 (std: 0.034122), ce qui représente la meilleure performance "
    "parmi les fonctions coût testées."
)

# IV – Optimizer (paragraph 79)
fill_question_run(
    paragraphs[79],
    "La méthode optimale est Adam car sur le test son mean_test_score est de "
    "0.712079 (std: 0.031625), ce qui est le meilleur score parmi tous les "
    "optimiseurs testés."
)

# V – Dropout / weight_constraint (paragraph 85)
fill_question_run(
    paragraphs[85],
    "La combinaison optimale est Dropout=0.0 et weight_constraint=3 car sur le test "
    "son mean_test_score est de 0.712000 (std: 0.032984), ce qui est le meilleur "
    "score parmi toutes les combinaisons testées."
)

# Construire un modèle optimal Exo2 (paragraph 92 – vide)
code_exo2 = (
    "from keras.models import Sequential\n"
    "from keras.layers import Dense, Dropout\n"
    "from keras.constraints import max_norm\n"
    "from scikeras.wrappers import KerasClassifier\n\n"
    "def create_optimal_model():\n"
    "    model = Sequential()\n"
    "    model.add(Dense(8, input_dim=8,\n"
    "                    kernel_initializer='lecun_uniform',\n"
    "                    activation='softplus',\n"
    "                    kernel_constraint=max_norm(3)))\n"
    "    model.add(Dropout(0.0))\n"
    "    model.add(Dense(1, activation='sigmoid'))\n"
    "    model.compile(loss='binary_crossentropy', optimizer='Adam',\n"
    "                  metrics=['accuracy'])\n"
    "    return model\n\n"
    "optimal_model = KerasClassifier(build_fn=create_optimal_model,\n"
    "                                epochs=100, batch_size=10, verbose=1)\n"
    "history = optimal_model.fit(X_train, y_train,\n"
    "                            validation_data=(X_test, y_test))\n"
    "# Résultats : Train_Loss=0.5650 | Train_Accuracy=0.7121\n"
    "#             Val_Loss=0.6339   | Val_Accuracy=0.7047"
)
add_text_to_empty_paragraph(paragraphs[92], code_exo2)

# Observations Exo2 (paragraph 94)
obs_exo2 = (
    "Le modèle optimal obtient une précision de 71.21% sur les données "
    "d'entraînement avec une perte de 0.5650. Sur les données de validation, "
    "la précision est de 70.47% avec une perte de 0.6339. L'écart entre les "
    "performances d'entraînement et de validation est très faible (< 1%), ce "
    "qui indique une excellente généralisation du modèle et l'absence "
    "d'overfitting significatif. L'initialisation lecun_uniform combinée à "
    "l'activation softplus et à l'optimiseur Adam s'avère particulièrement "
    "efficace pour ce problème de classification binaire sur la base Diabetes."
)
replace_placeholder_paragraph(paragraphs[94], obs_exo2)

# Conclusions Exo2 (paragraph 96)
conc_exo2 = (
    "La recherche par grille a permis d'identifier les hyperparamètres optimaux "
    "pour la classification binaire sur la base Diabetes : initialisation "
    "lecun_uniform, activation softplus, fonction coût binary_crossentropy, "
    "optimiseur Adam, dropout=0.0 et weight_constraint=3. Le modèle final atteint "
    "une précision de 71.21% en entraînement et 70.47% en validation, démontrant "
    "une excellente généralisation avec un écart minimal. Ces résultats confirment "
    "l'efficacité du tuning par GridSearchCV pour optimiser les hyperparamètres "
    "d'un ANN. Comparé à l'Exo1, ce modèle de classification présente une bien "
    "meilleure généralisation, probablement due à la nature du problème et à une "
    "meilleure adéquation des hyperparamètres sélectionnés."
)
replace_placeholder_paragraph(paragraphs[96], conc_exo2)

# ─────────────────────────────────────────────────────────────────────────────
doc.save('Devoir2_INFO0804_filled.docx')
print("Document sauvegardé avec succès.")
