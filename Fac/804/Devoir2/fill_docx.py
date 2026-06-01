import docx

doc = docx.Document('C:/Users/cocue/Documents/Fac/804/Devoir2/Devoir2_INFO0804.docx')

# ============================================================
# EXO 1 — HOUSING REGRESSION
# ============================================================

# Table 1: Exo1 I) Init poids
data_init_reg = [
    ('-33.252957', '5.765042'),
    ('-35.216032', '4.867743'),
    ('-497.553892', '28.662053'),
    ('-47.115650', '15.715966'),
    ('-36.011831', '5.989779'),
    ('-53.924160', '2.201705'),
    ('-43.486046', '15.910328'),
]
t = doc.tables[1]
for i, (mean, std) in enumerate(data_init_reg, start=1):
    t.rows[i].cells[1].text = mean
    t.rows[i].cells[2].text = std

# Table 2: Exo1 I) Optimal
t = doc.tables[2]
t.rows[1].cells[0].text = 'uniform'
t.rows[1].cells[1].text = '-33.252957'
t.rows[1].cells[2].text = '5.765042'

# Table 3: Exo1 II) Activation
data_act_reg = [
    ('-412.129668', '25.772618'),
    ('-28.958168',   '3.521216'),
    ('-92.382991',   '6.277295'),
    ('-38.397739',   '6.445264'),
    ('-85.811937',   '4.554786'),
    ('-105.501029',  '1.010475'),
    ('-135.833968', '13.521508'),
    ('-38.015256',   '3.386699'),
]
t = doc.tables[3]
for i, (mean, std) in enumerate(data_act_reg, start=1):
    t.rows[i].cells[1].text = mean
    t.rows[i].cells[2].text = std

# Table 4: Exo1 II) Optimal
t = doc.tables[4]
t.rows[1].cells[0].text = 'softplus'
t.rows[1].cells[1].text = '-28.958168'
t.rows[1].cells[2].text = '3.521216'

# Table 5: Exo1 III) Loss
data_loss_reg = [
    ('-25.258120',      '4.223791'),
    ('-33.796289',      '9.040761'),
    ('-244.475555',   '288.914290'),
    ('-50.187840',      '7.387986'),
    ('-12282.565797', '9918.070493'),
]
t = doc.tables[5]
for i, (mean, std) in enumerate(data_loss_reg, start=1):
    t.rows[i].cells[1].text = mean
    t.rows[i].cells[2].text = std

# Table 6: Exo1 III) Optimal
t = doc.tables[6]
t.rows[1].cells[0].text = 'mean_squared_error'
t.rows[1].cells[1].text = '-25.258120'
t.rows[1].cells[2].text = '4.223791'

# Table 7: Exo1 IV) Optimizer
data_opt_reg = [
    ('-88.225349',   '4.621445'),
    ('-30.021903',   '9.163522'),
    ('-87.601603',   '5.073488'),
    ('-257.510407', '189.391128'),
    ('-31.127648',   '1.153829'),
    ('-43.209253',   '7.453587'),
    ('-31.291924',   '5.141205'),
]
t = doc.tables[7]
for i, (mean, std) in enumerate(data_opt_reg, start=1):
    t.rows[i].cells[1].text = mean
    t.rows[i].cells[2].text = std

# Table 8: Exo1 IV) Optimal
t = doc.tables[8]
t.rows[1].cells[0].text = 'RMSprop'
t.rows[1].cells[1].text = '-30.021903'
t.rows[1].cells[2].text = '9.163522'

# Table 10: Exo1 V) Dropout
data_drop_reg = [
    ('-31.782132',   '3.355955'),
    ('-28.652179',   '3.866792'),
    ('-38.023699',   '9.148889'),
    ('-34.234945',   '5.627285'),
    ('-37.944128',   '8.145747'),
    ('-39.967745',   '9.354991'),
    ('-46.945604',  '18.200647'),
    ('-42.896364',  '10.546768'),
    ('-50.176487',  '23.029910'),
    ('-48.285971',   '8.716526'),
    ('-50.231848',   '9.272687'),
    ('-57.116230',   '3.784608'),
    ('-147.376980', '22.207695'),
    ('-174.013759', '25.676075'),
    ('-168.989367', '27.694494'),
]
t = doc.tables[10]
for i, (mean, std) in enumerate(data_drop_reg, start=1):
    t.rows[i].cells[1].text = mean
    t.rows[i].cells[2].text = std

# Table 11: Exo1 V) Optimal
t = doc.tables[11]
t.rows[1].cells[0].text = '{0.0, 2}'
t.rows[1].cells[1].text = '-28.652179'
t.rows[1].cells[2].text = '3.866792'

# Table 13: Exo1 Final Model (Train_Accuracy = MAE pour regression)
final_reg = [
    ('uniform',            '17.6971', '3.1349 (MAE)', '47.4620', '5.5789 (MAE)'),
    ('softplus',           '17.6971', '3.1349 (MAE)', '47.4620', '5.5789 (MAE)'),
    ('mean_squared_error', '17.6971', '3.1349 (MAE)', '47.4620', '5.5789 (MAE)'),
    ('RMSprop',            '17.6971', '3.1349 (MAE)', '47.4620', '5.5789 (MAE)'),
    ('0.0 / 2',            '17.6971', '3.1349 (MAE)', '47.4620', '5.5789 (MAE)'),
]
t = doc.tables[13]
for i, (val, tl, ta, vl, va) in enumerate(final_reg, start=1):
    t.rows[i].cells[1].text = val
    t.rows[i].cells[2].text = tl
    t.rows[i].cells[3].text = ta
    t.rows[i].cells[4].text = vl
    t.rows[i].cells[5].text = va

# ============================================================
# EXO 2 — DIABETES CLASSIFICATION
# ============================================================

# Table 14: Exo2 I) Init poids
data_init_cls = [
    ('0.719763', '0.029699'),
    ('0.719808', '0.047222'),
    ('0.657509', '0.051511'),
    ('0.671155', '0.036262'),
    ('0.671166', '0.045397'),
    ('0.634310', '0.036316'),
    ('0.688664', '0.022189'),
]
t = doc.tables[14]
for i, (mean, std) in enumerate(data_init_cls, start=1):
    t.rows[i].cells[1].text = mean
    t.rows[i].cells[2].text = std

# Table 15: Exo2 I) Optimal
t = doc.tables[15]
t.rows[1].cells[0].text = 'lecun_uniform'
t.rows[1].cells[1].text = '0.719808'
t.rows[1].cells[2].text = '0.047222'

# Table 16: Exo2 II) Activation
data_act_cls = [
    ('0.657543', '0.034938'),
    ('0.710028', '0.034506'),
    ('0.661453', '0.017007'),
    ('0.690614', '0.033481'),
    ('0.649757', '0.018080'),
    ('0.661408', '0.054199'),
    ('0.643864', '0.049102'),
    ('0.704202', '0.051440'),
]
t = doc.tables[16]
for i, (mean, std) in enumerate(data_act_cls, start=1):
    t.rows[i].cells[1].text = mean
    t.rows[i].cells[2].text = std

# Table 17: Exo2 II) Optimal
t = doc.tables[17]
t.rows[1].cells[0].text = 'softplus'
t.rows[1].cells[1].text = '0.710028'
t.rows[1].cells[2].text = '0.034506'

# Table 18: Exo2 III) Loss
data_loss_cls = [
    ('0.696405', '0.034122'),
    ('0.673127', '0.053312'),
    ('0.430210', '0.129794'),
]
t = doc.tables[18]
for i, (mean, std) in enumerate(data_loss_cls, start=1):
    t.rows[i].cells[1].text = mean
    t.rows[i].cells[2].text = std

# Table 19: Exo2 III) Optimal
t = doc.tables[19]
t.rows[1].cells[0].text = 'binary_crossentropy'
t.rows[1].cells[1].text = '0.696405'
t.rows[1].cells[2].text = '0.034122'

# Table 20: Exo2 IV) Optimizer
data_opt_cls = [
    ('0.673047', '0.039802'),
    ('0.675031', '0.043240'),
    ('0.519675', '0.122555'),
    ('0.485969', '0.152415'),
    ('0.712079', '0.031625'),
    ('0.661453', '0.010337'),
    ('0.696416', '0.039621'),
]
t = doc.tables[20]
for i, (mean, std) in enumerate(data_opt_cls, start=1):
    t.rows[i].cells[1].text = mean
    t.rows[i].cells[2].text = std

# Table 21: Exo2 IV) Optimal
t = doc.tables[21]
t.rows[1].cells[0].text = 'Adam'
t.rows[1].cells[1].text = '0.712079'
t.rows[1].cells[2].text = '0.031625'

# Table 23: Exo2 V) Dropout
data_drop_cls = [
    ('0.704180', '0.043624'),
    ('0.692518', '0.036477'),
    ('0.712000', '0.032984'),
    ('0.684755', '0.058421'),
    ('0.665273', '0.050246'),
    ('0.690625', '0.029395'),
    ('0.673081', '0.062762'),
    ('0.649712', '0.047617'),
    ('0.665307', '0.040509'),
    ('0.663357', '0.051664'),
    ('0.653588', '0.051664'),
    ('0.657509', '0.040341'),
    ('0.655549', '0.051071'),
    ('0.655560', '0.050339'),
    ('0.651673', '0.047609'),
]
t = doc.tables[23]
for i, (mean, std) in enumerate(data_drop_cls, start=1):
    t.rows[i].cells[1].text = mean
    t.rows[i].cells[2].text = std

# Table 24: Exo2 V) Optimal
t = doc.tables[24]
t.rows[1].cells[0].text = '{0.0, 3}'
t.rows[1].cells[1].text = '0.712000'
t.rows[1].cells[2].text = '0.032984'

# Table 26: Exo2 Final Model
final_cls = [
    ('lecun_uniform',       '0.5650', '0.7121', '0.6339', '0.7047'),
    ('softplus',            '0.5650', '0.7121', '0.6339', '0.7047'),
    ('binary_crossentropy', '0.5650', '0.7121', '0.6339', '0.7047'),
    ('Adam',                '0.5650', '0.7121', '0.6339', '0.7047'),
    ('0.0 / 3',             '0.5650', '0.7121', '0.6339', '0.7047'),
]
t = doc.tables[26]
for i, (val, tl, ta, vl, va) in enumerate(final_cls, start=1):
    t.rows[i].cells[1].text = val
    t.rows[i].cells[2].text = tl
    t.rows[i].cells[3].text = ta
    t.rows[i].cells[4].text = vl
    t.rows[i].cells[5].text = va

doc.save('C:/Users/cocue/Documents/Fac/804/Devoir2/Devoir2_INFO0804.docx')
print('Document rempli et sauvegarde avec succes.')
